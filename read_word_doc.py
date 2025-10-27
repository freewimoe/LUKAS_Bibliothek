"""
Script to read the content of the Word document
"""
from docx import Document

def read_word_document(file_path):
    """Read and print the content of a Word document"""
    try:
        doc = Document(file_path)
        
        print("=" * 80)
        print("CONTENT OF: LUKAS Bibliothek chatgpt.docx")
        print("=" * 80)
        print()
        
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():  # Only print non-empty paragraphs
                print(para.text)
                print()
        
        # Also check for tables if any
        if doc.tables:
            print("\n" + "=" * 80)
            print("TABLES FOUND IN DOCUMENT")
            print("=" * 80)
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\nTable {table_idx}:")
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    print(" | ".join(row_data))
    
    except Exception as e:
        print(f"Error reading document: {e}")
        print("\nTrying to install python-docx package...")
        import subprocess
        subprocess.run(["pip", "install", "python-docx"])
        print("\nPlease run the script again after installation.")

if __name__ == "__main__":
    file_path = r"c:\Users\fwmoe\Dropbox\Lukas_Wir-für-Lukas\Bibliothek\LUKAS Bibliothek chatgpt.docx"
    read_word_document(file_path)

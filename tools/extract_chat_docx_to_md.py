import sys
import os
import json
import hashlib
from datetime import datetime

try:
    from docx import Document
except Exception:
    print("ERROR: python-docx not installed. Please run: python -m pip install python-docx", file=sys.stderr)
    sys.exit(2)


def extract_docx_to_markdown(docx_path: str) -> str:
    doc = Document(docx_path)
    lines: list[str] = []

    # Simple header
    lines.append(f"## Import aus DOCX – {os.path.basename(docx_path)}")
    lines.append("")

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.rstrip()
        if text:
            lines.append(text)
        else:
            lines.append("")

    # Tables (if any): render as simple pipe-separated rows
    if doc.tables:
        lines.append("")
        lines.append("### Tabellen (aus DOCX)")
        for t_idx, table in enumerate(doc.tables, start=1):
            lines.append("")
            lines.append(f"Table {t_idx}:")
            for row in table.rows:
                row_data = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                lines.append(" | ".join(row_data))

    return "\n".join(lines) + "\n"


def append_to_file(path: str, content: str) -> None:
    # Ensure directory exists
    os.makedirs(os.path.dirname(path), exist_ok=True)
    # Append with UTF-8
    with open(path, "a", encoding="utf-8") as f:
        f.write(content)


def _state_path(md_path: str) -> str:
    # Store state alongside output md in the same folder
    folder = os.path.dirname(md_path) or "."
    return os.path.join(folder, ".chat_ingest_state.json")


def _load_state(state_path: str) -> dict:
    if os.path.exists(state_path):
        try:
            with open(state_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def _save_state(state_path: str, state: dict) -> None:
    try:
        with open(state_path, "w", encoding="utf-8") as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def main(argv: list[str]) -> int:
    if len(argv) < 3:
        print("Usage: extract_chat_docx_to_md.py <docx_path> <md_out_path>")
        return 1
    docx_path = argv[1]
    md_path = argv[2]

    if not os.path.exists(docx_path):
        print(f"ERROR: DOCX not found: {docx_path}", file=sys.stderr)
        return 2

    # Extract and compute content hash for deduplication
    md = extract_docx_to_markdown(docx_path)
    md_bytes = md.encode("utf-8")
    content_hash = hashlib.sha256(md_bytes).hexdigest()

    state_path = _state_path(md_path)
    state = _load_state(state_path)
    doc_key = os.path.abspath(docx_path)
    prev_hash = state.get(doc_key)

    if prev_hash == content_hash:
        print("No changes detected in DOCX; skip appending.")
        return 0

    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = [
        "\n",
        f"# RAW Chat Import – {ts}",
        "",
        "Hinweis: Import aus existierendem DOCX-Export. Falls nicht vollständig, fehlende Chat-Teile bitte manuell ergänzen.",
        "",
    ]

    append_to_file(md_path, "\n".join(header) + md)
    state[doc_key] = content_hash
    _save_state(state_path, state)
    print(f"Appended DOCX content to {md_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
